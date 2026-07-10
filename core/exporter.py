from docx import Document
from docx.shared import Pt, RGBColor
import csv
import json
from typing import List, Dict

def export_docx(leads: List[Dict], filename: str):
    """Exports generated leads list to a structured Word Document matching the sample format."""
    print(f"[DOCX] Exporting leads to DOCX: {filename}...")
    doc = Document()
    
    # Set default styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Title
    title = doc.add_heading('Detailed Prospect Report', level=0)
    title.runs[0].font.name = 'Calibri'
    title.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D) # Navy Blue
    
    intro = doc.add_paragraph(
        "Timidly Inc Priority Prospects — contact details pulled via Apify & Firecrawl, with "
        "tailored pitch angles based on the point of contact's background and company stage."
    )
    intro.runs[0].font.italic = True
    
    for idx, lead in enumerate(leads, 1):
        # Heading 1: e.g. "1. Autumn AI"
        h1 = doc.add_heading(f"{idx}. {lead['company_name']}", level=1)
        h1.runs[0].font.name = 'Calibri'
        h1.runs[0].font.size = Pt(14)
        h1.runs[0].font.bold = True
        h1.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
        
        # Tagline / Description
        p_tagline = doc.add_paragraph()
        tagline_text = lead.get('tagline', '')
        if tagline_text:
            run_tag = p_tagline.add_run(tagline_text)
            run_tag.font.italic = True
            run_tag.font.size = Pt(11)
        
        # Details list
        fields = [
            ("Point of Contact", lead.get("contact_name", "Not found")),
            ("Title", lead.get("contact_title", "GTM/Marketing Lead")),
            ("Email", lead.get("email", "Not found")),
            ("LinkedIn", lead.get("linkedin", "Not listed")),
            ("X / Twitter (X)", lead.get("twitter", "Not listed")),
            ("Mobile Phone (Number)", lead.get("phone", "Not found")),
            ("Country it is Based In", lead.get("country_based_in", "Unknown")),
            ("Funding Status", lead.get("funding", "Unknown / Seed")),
            ("Background of the Founders", lead.get("background_of_founders", "Not listed")),
            ("Lead Score", f"{lead.get('lead_score', '8')}/10"),
            ("Score Justification", lead.get("score_justification", "Not available")),
            ("Why This Company is a Lead", lead.get("why_pitch_fits", "")),
            ("Recommended Package", lead.get("recommended_package", "")),
            ("Tailored Outreach Angle", lead.get("tailored_outreach_angle", ""))
        ]
        
        for field_name, field_value in fields:
            p_field = doc.add_paragraph()
            run_name = p_field.add_run(f"{field_name}\n")
            run_name.bold = True
            run_name.font.size = Pt(10)
            run_name.font.color.rgb = RGBColor(0x59, 0x59, 0x59) # Gray
            
            run_val = p_field.add_run(f"{field_value}\n")
            run_val.font.size = Pt(11)
            
        doc.add_paragraph("---") # Section divider
        
    doc.save(filename)
    print("[OK] DOCX saved successfully.")

def export_csv(leads: List[Dict], filename: str):
    """Exports generated leads list to CSV for CRM import."""
    print(f"[CSV] Exporting leads to CSV: {filename}...")
    headers = [
        "company_name", "tagline", "contact_name", "contact_title", 
        "email", "linkedin", "twitter", "phone", "country_based_in", "funding", 
        "background_of_founders", "lead_score", "score_justification", "why_pitch_fits", 
        "recommended_package", "tailored_outreach_angle"
    ]
    
    with open(filename, mode='w', newline='', encoding='utf-8') as f:
        writer = csv.DictWriter(f, fieldnames=headers)
        writer.writeheader()
        for lead in leads:
            row = {h: lead.get(h, "Not found" if "email" in h or "phone" in h else "Not listed") for h in headers}
            writer.writerow(row)
    print("[OK] CSV saved successfully.")

def export_json(leads: List[Dict], filename: str):
    """Exports generated leads list to machine-readable JSON."""
    print(f"[JSON] Exporting leads to JSON: {filename}...")
    with open(filename, 'w', encoding='utf-8') as f:
        json.dump(leads, f, indent=4)
    print("[OK] JSON saved successfully.")
